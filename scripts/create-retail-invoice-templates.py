"""
Creates 4 retail invoice DOCX templates for KHCN disbursement module.
Uses [bracket] placeholder syntax matching the project's docxtemplater config.
Loop syntax: [#items]...[/items]
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"report_assets\KHCN templates\Chứng từ giải ngân"

def set_cell_text(cell, text, bold=False, italic=False, font_size=10, align=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def add_signature_table(doc):
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), "NGƯỜI MUA / BÊN B\n(Ký, ghi rõ họ tên)", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), "ĐẠI DIỆN BÊN BÁN\n(Ký tên, đóng dấu nếu có)", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 0), "\n\n", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 1), "\n\n", align=WD_ALIGN_PARAGRAPH.CENTER)

# ─── Template 1: Tạp hóa / Đồ uống ───────────────────────────────────────────

def create_mau1():
    doc = Document()
    # Supplier header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[supplier_name]")
    r.bold = True; r.font.size = Pt(13)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Địa chỉ: [supplier_address]    ĐT: [supplier_phone]").font.size = Pt(10)

    # Title
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("HOÁ ĐƠN BÁN LẺ")
    r3.bold = True; r3.font.size = Pt(14)

    # Invoice number & date (1-row table)
    meta = doc.add_table(rows=1, cols=2)
    meta.style = "Table Grid"
    set_cell_text(meta.cell(0, 0), "Số HĐ: [invoice_number]")
    set_cell_text(meta.cell(0, 1), "Ngày: [issue_date]", align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph("Người mua: [customer_name]")
    doc.add_paragraph("Địa chỉ: [customer_address]")

    # Items table
    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    headers = ["STT", "Tên hàng hoá", "ĐVT", "Số lượng", "Đơn giá (đ)", "Thành tiền (đ)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Loop row — docxtemplater [#items]...[/items]
    loop_row = table.rows[1]
    loop_texts = ["[#items][i]", "[name]", "[unit]", "[qty]", "[unit_price_fmt]", "[subtotal_fmt][/items]"]
    for i, t in enumerate(loop_texts):
        set_cell_text(loop_row.cells[i], t, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_col_widths(table, [1.0, 6.0, 1.5, 1.5, 2.5, 2.5])

    doc.add_paragraph("Tổng cộng: [total_fmt] đồng")
    doc.add_paragraph("Bằng chữ: [total_words]")
    doc.add_paragraph()
    add_signature_table(doc)

    path = os.path.join(OUTPUT_DIR, "HoaDon_TapHoa_DoUong.docx")
    doc.save(path)
    print(f"Created: {path}")

# ─── Template 2: Vật liệu xây dựng ───────────────────────────────────────────

def create_mau2():
    doc = Document()
    # Company info table
    info = doc.add_table(rows=1, cols=2)
    info.style = "Table Grid"
    set_cell_text(info.cell(0, 0),
        "CÔNG TY / CỬA HÀNG [supplier_name]\nĐịa chỉ: [supplier_address]\nĐT: [supplier_phone]")
    set_cell_text(info.cell(0, 1),
        "HOÁ ĐƠN BÁN LẺ\nSố: [invoice_number]\nNgày: [issue_date]\nHình thức TT: [payment_method]",
        align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph("Khách hàng: [customer_name]")
    doc.add_paragraph("Địa chỉ: [customer_address]")

    table = doc.add_table(rows=2, cols=7)
    table.style = "Table Grid"
    headers = ["STT", "Tên hàng hoá / Vật tư", "ĐVT", "SL", "Đơn giá (đ)", "Thành tiền (đ)", "Ghi chú"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    loop_row = table.rows[1]
    loop_texts = ["[#items][i]", "[name]", "[unit]", "[qty]", "[unit_price_fmt]", "[subtotal_fmt]", "[note][/items]"]
    for i, t in enumerate(loop_texts):
        set_cell_text(loop_row.cells[i], t, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_col_widths(table, [1.0, 5.5, 1.5, 1.2, 2.0, 2.3, 2.0])

    doc.add_paragraph("Tổng tiền bằng chữ: [total_words]")
    doc.add_paragraph("Ghi chú: Hàng giao tại công trình, chưa bao gồm chi phí vận chuyển.")
    doc.add_paragraph()
    add_signature_table(doc)

    path = os.path.join(OUTPUT_DIR, "HoaDon_VatLieuXayDung.docx")
    doc.save(path)
    print(f"Created: {path}")

# ─── Template 3: Thiết bị y tế ────────────────────────────────────────────────

def create_mau3():
    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập – Tự do – Hạnh phúc").font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("HOÁ ĐƠN BÁN LẺ TRANG THIẾT BỊ Y TẾ")
    r2.bold = True; r2.font.size = Pt(13)

    info = doc.add_table(rows=1, cols=2)
    info.style = "Table Grid"
    set_cell_text(info.cell(0, 0),
        "BÊN BÁN (Bên A):\nTên: [supplier_name]\nĐịa chỉ: [supplier_address]\nSĐT: [supplier_phone]")
    set_cell_text(info.cell(0, 1),
        "Số HĐ: [invoice_number]\nNgày lập: [issue_date]\n\nBÊN MUA (Bên B):\nHọ tên: [customer_name]\nĐịa chỉ: [customer_address]")

    table = doc.add_table(rows=2, cols=7)
    table.style = "Table Grid"
    headers = ["STT", "Tên trang thiết bị / Vật tư y tế", "ĐVT", "SL", "Đơn giá (đ)", "Thành tiền (đ)", "Ghi chú"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    loop_row = table.rows[1]
    loop_texts = ["[#items][i]", "[name]", "[unit]", "[qty]", "[unit_price_fmt]", "[subtotal_fmt]", "[note][/items]"]
    for i, t in enumerate(loop_texts):
        set_cell_text(loop_row.cells[i], t, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_col_widths(table, [1.0, 5.5, 1.5, 1.2, 2.0, 2.3, 2.0])

    doc.add_paragraph("Tổng cộng: [total_fmt] đồng")
    doc.add_paragraph("Bằng chữ: [total_words]")
    doc.add_paragraph("Hình thức thanh toán: Tiền mặt")
    doc.add_paragraph("* Lưu ý: Hóa đơn này không có giá trị khấu trừ thuế GTGT.")
    doc.add_paragraph()
    add_signature_table(doc)

    path = os.path.join(OUTPUT_DIR, "HoaDon_ThietBiYTe.docx")
    doc.save(path)
    print(f"Created: {path}")

# ─── Template 4: Nông sản ─────────────────────────────────────────────────────

def create_mau4():
    doc = Document()
    info = doc.add_table(rows=1, cols=1)
    info.style = "Table Grid"
    set_cell_text(info.cell(0, 0),
        "HỘ KINH DOANH / CƠ SỞ: [supplier_name]\nĐịa chỉ: [supplier_address]    SĐT: [supplier_phone]",
        bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PHIẾU BÁN HÀNG / BIÊN NHẬN TIỀN")
    r.bold = True; r.font.size = Pt(13)

    doc.add_paragraph("Số: [invoice_number]        Ngày: [issue_date]")
    doc.add_paragraph("Người mua: [customer_name]")
    doc.add_paragraph("Địa chỉ: [customer_address]")

    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    headers = ["STT", "Tên nông sản / Hàng hoá", "ĐVT", "SL", "Đơn giá (đ)", "Thành tiền (đ)"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    loop_row = table.rows[1]
    loop_texts = ["[#items][i]", "[name]", "[unit]", "[qty]", "[unit_price_fmt]", "[subtotal_fmt][/items]"]
    for i, t in enumerate(loop_texts):
        set_cell_text(loop_row.cells[i], t, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_col_widths(table, [1.0, 6.0, 1.5, 1.5, 2.5, 2.5])

    doc.add_paragraph("Bằng chữ: [total_words]")
    doc.add_paragraph("Thanh toán: [payment_method]")
    doc.add_paragraph("Hàng đã bán không nhận lại. Vui lòng kiểm tra hàng trước khi nhận./.")
    doc.add_paragraph()
    add_signature_table(doc)

    path = os.path.join(OUTPUT_DIR, "HoaDon_NongSan.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_mau1()
    create_mau2()
    create_mau3()
    create_mau4()
    print("All 4 templates created.")
